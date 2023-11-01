from .models import Analisis, Aplicacion, Resultado, Modelo, Carpeta, Archivo, Preferencias, Grafico, Grafico_Imagen, Tabla
import json
import spacy
from spacy import displacy
import os
import pandas as pd
import altair as alt
from altair_saver import save
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from django.core.files import File 
import io
import re
import datetime
import docx

def procesar_analisis(analisis, user):
    """ 
    Esta funcion realiza el procesamiento pedido por el usuario y genera los resultados.
    Por ahora la funcion solo permite aplicar modelos de entidades y clasificadores.
    Pero se podria extender para que permita aplicar otros modelos de NLP.
    
    Parameters
    ----------

    :param: analisis (Analisis) El analisis que recibe tiene: informe vacio, fecha actual, carpeta y modelo elegidos.
    :param: user (User) El usuario que realiza el analisis.

    """
    carpeta = analisis.carpeta
    modelo = analisis.modelo 

    if modelo.nombre == 'entidades': #cargo ese solo modelo y lo aplico
        procesar_entidades(analisis, carpeta, user)
    elif modelo.nombre == 'clasificador': # aplico primero el modelo binario y despues el modelo multicategoria.  
        procesar_clasificador(analisis, carpeta, user)
    else:
        raise FileNotFoundError("El modelo no existe") #podria cambiar esto, intentar cargar el modelo y si no existe tirar un error. Y pasar a hacer un procesamiento 'generico' que no dependa de un modelo en particular.

def procesar_entidades(analisis, carpeta, user):
    """ 
    Esta función se encarga de procesar los archivos de una carpeta con el modelo de detección de entidades.
    Para cada archivo, se extraen los textos y se detectan las entidades.
    Luego, se arma el objeto resultado y se lo guarda en la base de datos.
    Finalmente, se arma el informe con los resultados y se lo guarda en el analisis.

    Parameters
    ----------
    :param: analisis (Analisis) El analisis que recibe tiene: informe vacio, fecha actual, carpeta y modelo elegidos.
    :param: carpeta (Carpeta) La carpeta que recibe tiene: nombre, usuario y archivos.
    :param: user (User) El usuario que realiza el analisis.

    Return
    ------
    :return: 1 (int) Si el analisis se realizo correctamente devuelve 1.
    """
    
    # 1: Cargar el modelo de spacy
    model_path = os.path.join(os.getcwd(),"analisis", "static", "modelos", "entidades" )
    nlp = spacy.load(model_path)

    colors_azul_amarillo = {"DINERO": "#00065D", 
                            "FECHA":"#2706AD", 
                            "HORA": "#5C40D0", 
                            "LUGAR": "#889AFD",
                            "MEDIDA": "#CCC5B2",
                            "MISC": "#F3D850",
                            "ORG": "#FFD500",
                            "PERSONA": "#9A7300",
                            "TIEMPO" : "#DFA000",
                            }
    
    colors_azul_rosa = {"DINERO": "#00065D", 
                            "FECHA":"#2257EC", 
                            "HORA": "#008DDE", 
                            "LUGAR": "#7A8BE6",
                            "MEDIDA": "#D2C3C3",
                            "MISC": "#FF9696",
                            "ORG": "#F94D67",
                            "PERSONA": "#E9255A",
                            "TIEMPO" : "#99001C",
                            }
    colors_default = {"DINERO": "#F3D850", 
                            "FECHA":"#FECA74", 
                            "HORA": "#7AECEC", 
                            "LUGAR": "#BFE1D9",
                            "MEDIDA": "#AA9CFC",
                            "MISC": "#C887FB",
                            "ORG": "#E4E7D2",
                            "PERSONA": "#FFB700",
                            "TIEMPO" : "#52CCCC",
                            }

    preferencia = Preferencias.objects.get(usuario = user)
    if preferencia.color == "AM":
        options = {"colors": colors_azul_amarillo}
    elif preferencia.color == "AR":
        options = {"colors": colors_azul_rosa}
    else:
        options = {"colors": colors_default}

    # 2: Extraer textos de los archivos en la carpeta
    archivos = Archivo.objects.filter(carpeta = carpeta)
    for archivo in archivos:
        lines = []
        nombre, partition, extension = archivo.arch.name.rpartition('.')
        if extension == 'docx':
            doc = docx.Document(archivo.arch.path)
            for i in doc.paragraphs:
                print(i.text)
                lines.append(i.text)
        elif extension == 'xlsx':
            df = pd.read_excel(archivo.arch.path, usecols="A") #siempre leo la primera columna
            lines = df.values.tolist()
        elif extension == 'txt':
            with open(archivo.arch.path, "r", encoding = 'UTF-8') as f:
                lines = f.read().splitlines()
        else: 
            #file not supported exception
            #avisar al usuario que la carpeta tiene un archivo en un formato no soportado
            raise ValueError(f'Tipo de archivo no soportado: {extension}', archivo)
        docs = list(nlp.pipe(lines))

        # 3: Armar el objeto resultado de cada uno:
        for index, doc in enumerate(docs):
            entidades = []
            for ent in doc.ents:
                entidad = {
                    "text": ent.text,
                    "label": ent.label_
                }
                entidades.append(entidad)
            entidades_json = json.dumps(entidades, ensure_ascii=False)
            texto = doc.text
            detectado = entidades_json
            html = displacy.render(doc, style="ent", options=options)
            archivo_origen = archivo
            numero_linea = index + 1
            Resultado(texto= texto, detectado = detectado, html = html, numero_linea = numero_linea, analisis = analisis, archivo_origen = archivo_origen).save()

    # 4: Procesar los resultados y armar el informe segun el modelo que sea
    armar_informe_entidades(analisis, preferencia)

    return 1

def armar_informe_entidades(analisis, preferencia): 
    """
    Esta funcion crea los graficos y las tablas que contiene el informe de los resultados de un analisis creado por un modelo de deteccion de entidades.

    Parameters
    ----------
    :param: analisis (Analisis) Es un analisis realizado previamente que ya tiene resultados asociados, de los cuales se quiere crear un informe.

    :param: preferencia (Preferencias) Es un objeto que contiene las preferencias del usuario que realizo el analisis.

    Return
    ------
    :return: 1 (int) Si los elementos del informe se crearon correctamente devuelve 1.

    """

    os.makedirs(f'analisis/static/graficos/{analisis.id}', exist_ok=True)

    resultados = Resultado.objects.filter(analisis = analisis)
    
    domain =["DINERO", "FECHA", "HORA", "LUGAR", "MEDIDA", "MISC", "ORG", "PERSONA", "TIEMPO"]
    range_ = []

    if preferencia.color == "AM":
        range_= ["#00065D","#2706AD","#5C40D0","#889AFD", "#CCC5B2","#F3D850","#FFD500","#9A7300", "#DFA000"]
    elif preferencia.color == "AR":
        range_= ["#00065D", "#2257EC","#008DDE","#7A8BE6","#D2C3C3","#FF9696","#F94D67","#E9255A","#99001C"]
    else:
        range_ = ["#F3D850","#FECA74", "#7AECEC", "#BFE1D9", "#AA9CFC", "#C887FB" ,"#E4E7D2" ,"#FFB700", "#52CCCC"]

    
    data = []
    for resultado in resultados:
        entidades = json.loads(resultado.detectado)

        for entidad in entidades:
            data.append({
                'text': entidad['text'],
                'label': entidad['label'],
                'archivo_origen': resultado.archivo_origen.nombre,
                'numero_linea': resultado.numero_linea
            })
    df = pd.DataFrame(data)

    total_entidades = df.shape[0]
    entidades_counts = df['label'].value_counts().reset_index()#pd indice tipos de entidades y valor num de apariciones
    entidades_counts.columns = ['Etiqueta', 'Apariciones']
    tabla_entidades_count = Tabla(nombre = "Distribucion de entidades", tabla = entidades_counts.to_html(classes='table table-striped table-hover table-sm', index=False), analisis = analisis)
    tabla_entidades_count.save()



    ent_text_counts = df['text'].value_counts() #entidades que se repitan: tendrian el count en mas de 1
    repeating_entities = ent_text_counts[ent_text_counts > 1] #entidades que se repiten
    num_rep = repeating_entities.shape[0] #numero de entidades que se repiten
    if num_rep > 0:
        tabla_rep_ents = repeating_entities.reset_index()
        tabla_rep_ents.columns = ['Entidad', 'Apariciones'] 
        tabla_repeating_ents = Tabla(nombre = "Entidades que se repiten", tabla = tabla_rep_ents.to_html(classes='table table-striped table-hover table-sm', index=False), analisis = analisis)
        tabla_repeating_ents.save()


    # Cantidad de cada tipo de entidades
    chart_count_ents = alt.Chart(df, title="Distribucion de entidades").mark_bar().encode(
        x = alt.X('label:N', title='Entidades'),
        y = alt.Y('count():Q', title='N° Apariciones'),
        color = alt.Color('label', scale = alt.Scale(domain=domain, range=range_)),
        tooltip=['label:N', 'count():Q']
    ).interactive()
    json_count_ents = chart_count_ents.to_json()
    grafico_cout_ents = Grafico(nombre = "Distribucion de entidades", chart = json_count_ents, analisis = analisis)
    grafico_cout_ents.save()


    chart_torta = alt.Chart(df).mark_arc().encode(
    theta="count():Q",
    color= alt.Color('label', scale = alt.Scale(domain=domain, range=range_)),
    tooltip=['label:N', 'count():Q']
    ).interactive()
    json_torta = chart_torta.to_json()
    grafico_torta = Grafico(nombre = "Torta distribucion de entidades", chart = json_torta, analisis = analisis)
    grafico_torta.save()

    #Entidades por cada archivo
    file_entity_counts = df.groupby('archivo_origen')['text'].count().reset_index()
    file_entity_counts.columns = ['archivo_origen', 'entity_count']
    chart_ents_file = alt.Chart(file_entity_counts).mark_bar().encode(
        x=alt.X('archivo_origen:N', title='Archivo'),
        y=alt.Y('entity_count:Q', title='Numero de entidades'),
        color=alt.Color('archivo_origen:N', title='File', scale = alt.Scale(range=range_)),
        tooltip=['archivo_origen:N', 'entity_count:Q']
    ).properties(
        title='Entidades por archivo'
    ).interactive()
    json_ents_file = chart_ents_file.to_json()
    grafico_ents_file = Grafico(nombre = "Entidades por archivo", chart = json_ents_file, analisis = analisis)
    grafico_ents_file.save()


    #Tipo entidades por cada archivo
    # Group and count entity labels by file
    file_label_counts = df.groupby(['archivo_origen', 'label'])['text'].count().reset_index()
    file_label_counts.columns = ['archivo_origen', 'label', 'count']
    # Create a stacked bar chart
    chart_entscomp_file = alt.Chart(file_label_counts).mark_bar().encode(
        x=alt.X('archivo_origen:N', title='File'),
        y=alt.Y('count:Q', title='Entity Count'),
        color=alt.Color('label:N', title='Entity Label',scale = alt.Scale(domain=domain, range=range_)),
        tooltip=['archivo_origen:N', 'label:N', 'count:Q']
    ).properties(
        title='Composicion de entidades por archivo'
    )
    json_entscomp_file = chart_entscomp_file.to_json()
    grafico_entscomp_file = Grafico(nombre = "Composicion de entidades por archivo", chart = json_entscomp_file, analisis = analisis)
    grafico_entscomp_file.save()

    num_ticks = 10
    ## Relacion para cada archivo: numero de entidades y linea 
    scatterplots = []
    for file_name, group_df in df.groupby('archivo_origen'):
        vector_valores = []
        max_v = group_df['numero_linea'].max()
        min_v = group_df['numero_linea'].min()
        range_v = max_v - min_v
        interval = max(1, int(range_v / num_ticks))
        current_value = min_v
        while current_value <= max_v:
                vector_valores.append(current_value)
                current_value += interval
        axis=alt.Axis(values=vector_valores)

        scatterplot = alt.Chart(group_df, width = 900).mark_circle().encode(
            x=alt.X('numero_linea:O', title='Numero de lines', axis=axis),
            y=alt.Y('label:N', title='Entidad'),
            color=alt.Color('label:N', title='Entidad', scale = alt.Scale(domain=domain, range=range_)),
            tooltip=['label:N', 'numero_linea:O', 'text:N']
        ).properties(
            title=f'Relación entre numero de linea y entidades en archivo {file_name}'
        )
        scatterplots.append(scatterplot)
    chart_lineas_entidades = alt.vconcat(*scatterplots)
    json_lineas_entidades = chart_lineas_entidades.to_json()
    grafico_lineas_entidades = Grafico(nombre = "Relacion entre numero de linea y entidades", chart = json_lineas_entidades, analisis = analisis)
    grafico_lineas_entidades.save()
    

    #Wordcloud de los textos de las entidades
    text = ' '.join(df['text'])
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title('Entidades detectadas')

    wordcloud_path = f'analisis/static/graficos/{analisis.id}/wordcloud_{analisis.id}.png'
    wordcloud.to_file(wordcloud_path)
    imagen_wordcloud_total = Grafico_Imagen(nombre = "Wordcloud de entidades", analisis = analisis)
    imagen_wordcloud_total.imagen.save(wordcloud_path, File(open(wordcloud_path, 'rb')))
    imagen_wordcloud_total.save()
    
    """
    # Wordcloud por cada tipo de entidad
    
    text_by_label = df.groupby('label')['text'].apply(lambda x: ' '.join(x))
    for label, text in text_by_label.items():
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)

        plt.figure(figsize=(8, 4))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.title(f'Word Cloud for Entity Label: {label}')
        
        wordcloud_path = f'analisis/static/graficos/{analisis.id}/wordcloud_{label}{analisis.id}.png'
        wordcloud.to_file(wordcloud_path)

        imagen_wordcloud = Grafico_Imagen(nombre = f"Wordcloud de entidades {label}", analisis = analisis)
        imagen_wordcloud.imagen.save(wordcloud_path, File(open(wordcloud_path, 'rb')))
        imagen_wordcloud.save()
    """

    return 1

def wpp_format(s):
    """
    Verificar si la línea tiene el formato de mensaje de WhatsApp.

    Parameters
    ----------
    :param: s (str) Línea del archivo.

    Returns
    -------
    :return: (bool) True si la línea tiene el formato de mensaje de WhatsApp, False en caso contrario.
    """
    pattern = r"\d{1,2}/\d{1,2}/\d{4}, \d{2}:\d{2} - .+: .+"
    match = re.match(pattern, s)
    return match is not None


def procesar_linea(analisis, archivo, line, index):
    elemento = {}
    if wpp_format(line):
        date_name_text = line.strip().split(' - ') # [date, name: text]
        name_index = date_name_text[1].find(':') 
        name = date_name_text[1][:name_index] 
        text = date_name_text[1][name_index+2:]
        date = date_name_text[0]
        fecha = datetime.datetime.strptime(date, "%d/%m/%Y, %H:%M").date()
        #Problema: fechas con dia y mes intercambiados (puedo hacer que si no esta en ese formato queda 'desconocida' y listo)

        if(text.find('(archivo adjunto)') != -1):
            texto = text.replace('(archivo adjunto)', '') 
            numero_linea = index + 1
            detectado = 'Adjunto'
            Resultado(texto= texto, detectado = detectado, html = "", numero_linea = numero_linea, analisis = analisis, archivo_origen = archivo, fecha_envio = fecha, remitente = name).save() 
        elif (text.find('<Multimedia omitido>') != -1):
            texto = "Adjunto Desconocido"
            detectado = 'Adjunto'
            numero_linea = elemento['numero_linea']
            Resultado(texto= texto, detectado = detectado, html = "", numero_linea = numero_linea, analisis = analisis, archivo_origen = archivo, fecha_envio = fecha, remitente = name).save()
        else:
            elemento = {
                'texto': text,
                'fecha': fecha,
                'remitente': name,
                'numero_linea': index + 1,
                'archivo_origen': archivo,
                'doc' : None
            }
    else:
        elemento = {
            'texto': line,
            'fecha': None,
            'remitente': None,
            'numero_linea': index + 1,
            'archivo_origen': archivo,
            'doc' : None
        }
    return elemento


def procesar_clasificador(analisis, carpeta, user):
    """
    Esta función se encarga de procesar los archivos de una carpeta con el modelo de clasificación de violencia. Crea los objetos Resultado en la base de datos para cada linea que se pidió analizar.

    Parameters
    ----------
    :param: analisis (Analisis) El analisis que recibe tiene: informe vacio, fecha actual, carpeta y modelo elegidos.
    :param: carpeta (Carpeta) La carpeta que recibe tiene: nombre, usuario y archivos.
    :param: user (User) El usuario que realiza el analisis.

    Return
    ------
    :return: 1 (int) Si el analisis se realizo correctamente devuelve 1.

    """
    
    model_path = os.path.join(os.getcwd(),"analisis", "static", "modelos", "clasificador-binario" )
    nlp = spacy.load(model_path)
    model_path_multi = os.path.join(os.getcwd(),"analisis", "static", "modelos", "clasificador-multi" )
    nlp_multi = spacy.load(model_path_multi)

    archivos = Archivo.objects.filter(carpeta = carpeta)
    for archivo in archivos:
        dict_text = []
        lines = []
        if archivo.arch.name.split('.')[1] == 'docx':
            doc = docx.Document(archivo.arch.path)
            for index, i in enumerate(doc.paragraphs):
                line = i.text
                elemento = procesar_linea(analisis= analisis, archivo= archivo, line= line, index= index)
                if elemento:
                    dict_text.append(elemento)
        else:        
            with open(archivo.arch.path, "r", encoding = 'UTF -8') as f:
                for index, line in enumerate(f):
                    elemento = procesar_linea(analisis= analisis, archivo= archivo, line= line, index= index)
                    if elemento:
                        dict_text.append(elemento)

        lines = [elemento['texto'] for elemento in dict_text] # tengo que sacar la comprobacion de si es o no adjunto porque necesito que los docs_binarios sean paralelos a dict_text no me los puedo saltear
        docs_binarios = list(nlp.pipe(lines))
        for index, doc in enumerate(docs_binarios):
            dict_text[index]['doc'] = doc

        violentos_text = []
        elementos_violentos = []
        for elemento in dict_text:
            doc = elemento['doc']
            if doc.cats['Violento'] < 0.8:
                texto = doc.text
                detectado = 'No Violento'
                html = ""
                archivo_origen = archivo
                numero_linea = elemento['numero_linea']
                fecha = elemento['fecha']
                remitente = elemento['remitente']
                Resultado(texto= texto, detectado = detectado, html = html, numero_linea = numero_linea, analisis = analisis, archivo_origen = archivo_origen, fecha_envio = fecha, remitente = remitente).save()
            else:
                elementos_violentos.append(elemento)

        violentos_text = [elemento['texto'] for elemento in elementos_violentos]
        docs_multi = list(nlp_multi.pipe(violentos_text))
        for index, doc in enumerate(docs_multi):
            elementos_violentos[index]['doc'] = doc

        for elemento in elementos_violentos:
            doc = elemento['doc']
            texto = doc.text
            detectado = max(doc.cats, key=doc.cats.get)
            html = ""
            archivo_origen = archivo
            numero_linea = elemento['numero_linea']
            fecha = elemento['fecha']
            remitente = elemento['remitente']
            Resultado(texto= texto, detectado = detectado, html = html, numero_linea = numero_linea, analisis = analisis, archivo_origen = archivo_origen, fecha_envio = fecha, remitente = remitente).save()


    # 4: Procesar los resultados y armar el informe segun el modelo que sea
    preferencia = Preferencias.objects.get(usuario = user)

    armar_informe_clasificador(analisis,preferencia)
    return 1
    


def armar_informe_clasificador(analisis, preferencia):
    """
    Esta funcion crea los graficos y las tablas que contiene el informe de los resultados de un analisis creado por un modelo de clasificación de violencia

    Parameters
    ----------
    :param: analisis (Analisis) Es un analisis realizado previamente que ya tiene resultados asociados, de los cuales se quiere crear un informe.

    :param: preferencia (Preferencias) Es un objeto que contiene las preferencias del usuario que realizo el analisis.

    Return
    ------
    :return: 1 (int) Si los elementos del informe se crearon correctamente devuelve 1.

    """

    os.makedirs(f'analisis/static/graficos/{analisis.id}', exist_ok=True)

    resultados = Resultado.objects.filter(analisis = analisis)
    
    domain =["No Violento", "Sexual", "Física", "Económica", "Simbólica", "Psicológica", "Adjunto"]
    domain_v =["Sexual", "Física", "Económica", "Simbólica", "Psicológica"]
    range_ = []

    if preferencia.color == "AM":
        range_= ["#2706AD","#5C40D0","#889AFD","#F3D850","#FFD500","#9A7300"]
    elif preferencia.color == "AR":
        range_= ["#00065D","#2257EC","#7A8BE6","#FF9696","#F94D67","#99001C"]
    else:
        range_= ["#F3D850","#FECA74","#7AECEC","#BFE1D9","#AA9CFC","#C887FB"]

    data_completos = []
    data_violentos = []
    for resultado in resultados:
        data_completos.append({
            'text': resultado.texto,
            'clasificacion': resultado.detectado,
            'archivo_origen': resultado.archivo_origen.nombre,
            'numero_linea': resultado.numero_linea
        })
        if (str(resultado.detectado) != "No Violento"):
                    data_violentos.append({
                        'text': resultado.texto,
                        'clasificacion': str(resultado.detectado),
                        'archivo_origen': resultado.archivo_origen.nombre,
                        'numero_linea': resultado.numero_linea
                    })

        

    df_completo = pd.DataFrame(data_completos)
    df_violentos = pd.DataFrame(data_violentos)

    counts = df_completo['clasificacion'].value_counts().reset_index()#pd indice categorias y valor num de apariciones
    counts.columns = ['Categoria', 'Apariciones']
    tabla_categorias_count = Tabla(nombre = "Distribucion de categorias", tabla = counts.to_html(classes='table table-striped table-hover table-sm', index=False), analisis = analisis)
    tabla_categorias_count.save()

    # -deberia excluir los que tienen valor No violento

    # Cantidad de cada tipo de violencia
    chart_count_cats = alt.Chart(df_violentos, title="Distribucion de categorias").mark_bar().encode(
        x = alt.X('clasificacion:N', title='Categoria'),
        y = alt.Y('count():Q', title='N° Apariciones'),
        color = alt.Color('clasificacion', scale = alt.Scale(domain=domain_v, range=range_)),
        tooltip=['clasificacion:N', 'count():Q']
    ).interactive()
    json_count_cats = chart_count_cats.to_json()
    grafico_cout_cats = Grafico(nombre = "Distribucion de categorias", chart = json_count_cats, analisis = analisis)
    grafico_cout_cats.save()

    chart_torta = alt.Chart(df_violentos).mark_arc().encode(
    theta="count():Q",
    color= alt.Color('clasificacion', scale = alt.Scale(domain=domain_v, range=range_)),
    tooltip=['clasificacion:N', 'count():Q']
    ).interactive()
    json_torta = chart_torta.to_json()
    grafico_torta = Grafico(nombre = "Torta distribucion de categorias", chart = json_torta, analisis = analisis)
    grafico_torta.save()

    #Tipo de violencia por cada archivo
    file_label_counts = df_violentos.groupby(['archivo_origen', 'clasificacion'])['text'].count().reset_index()
    file_label_counts.columns = ['archivo_origen', 'clasificacion', 'count']
    # Create a stacked bar chart
    chart_comp_file = alt.Chart(file_label_counts).mark_bar().encode(
        x=alt.X('archivo_origen:N', title='File'),
        y=alt.Y('count:Q', title='Categoria Count'),
        color=alt.Color('clasificacion:N', title='Clasificacion violencia',scale = alt.Scale(domain=domain_v, range=range_)),
        tooltip=['archivo_origen:N', 'clasificacion:N', 'count:Q']
    ).properties(
        title='Composicion de tipos de violencia por archivo'
    )
    json_comp_file = chart_comp_file.to_json()
    grafico_comp_file = Grafico(nombre = "Composicion de categorias por archivo", chart = json_comp_file, analisis = analisis)
    grafico_comp_file.save()

    
    ## Relacion para cada archivo: numero de violentos y linea 
    scatterplots = []
    num_ticks = 10
    for file_name, group_df in df_violentos.groupby('archivo_origen'):
        vector_valores = []
        max_v = group_df['numero_linea'].max()
        min_v = group_df['numero_linea'].min()
        range_v = max_v - min_v
        interval = max(1, int(range_v / num_ticks))
        current_value = min_v
        while current_value <= max_v:
                vector_valores.append(current_value)
                current_value += interval
        axis=alt.Axis(values=vector_valores)
        scatterplot = alt.Chart(group_df, width=900).mark_circle().encode(
            x=alt.X('numero_linea:O', title='Numero de lines', axis=axis),
            y=alt.Y('clasificacion:N', title='Categoria'),
            color=alt.Color('clasificacion:N', title='Categoria', scale = alt.Scale(domain=domain_v, range=range_)),
            tooltip=['clasificacion:N', 'numero_linea:O', 'text:N']
        ).properties(
            title=f'Relación entre numero de linea y frases violentas en archivo {file_name}'
        )
        scatterplots.append(scatterplot)
    chart_lineas_cat = alt.vconcat(*scatterplots)

    json_lineas_cat = chart_lineas_cat.to_json()
    grafico_lineas_cat = Grafico(nombre = "Relacion numero de linea y frases violentas", chart = json_lineas_cat, analisis = analisis)
    grafico_lineas_cat.save()

    #Wordcloud completo
    text = ' '.join(df_completo['text'])
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title('Palabras detectadas')
    wordcloud_path = f'analisis/static/graficos/{analisis.id}/wordcloud_cats_{analisis.id}.png'
    wordcloud.to_file(wordcloud_path)
    imagen_wordcloud_total = Grafico_Imagen(nombre = "Wordcloud de clasificacion", analisis = analisis)
    imagen_wordcloud_total.imagen.save(wordcloud_path, File(open(wordcloud_path, 'rb')))
    imagen_wordcloud_total.save()

    #Wordcloud violentos
    df_v = df_violentos.loc[df_violentos['clasificacion'] != 'Adjunto'] 
    if df_v.shape[0] > 0:
        text = ' '.join(df_v['text'])
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.title('Palabras detectadas')
        wordcloud_path = f'analisis/static/graficos/{analisis.id}/wordcloud_violentos_{analisis.id}.png'
        wordcloud.to_file(wordcloud_path)
        imagen_wordcloud_total = Grafico_Imagen(nombre = "Wordcloud de violentos", analisis = analisis)
        imagen_wordcloud_total.imagen.save(wordcloud_path, File(open(wordcloud_path, 'rb')))
        imagen_wordcloud_total.save()
    return 1